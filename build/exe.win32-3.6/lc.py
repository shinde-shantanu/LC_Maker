from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from docx import *
import tempfile
import win32api
import win32print


def give_print(dat):
    filename = 'lc - Copy.docx'
    #f=open(filename,'x')
    d=docx.Document(filename)
    l=d.paragraphs
    t=d.tables

##=df['अ. क्र. '][x]
##            =df['जनरल रजिस्टर क्र.'][x]
##            =df['स्टुडंट आय डी'][x]
##            =df['यु आय डी नं.'][x]
##            =df['विद्यार्थ्याचे संपूर्ण नाव'][x]
##            =df['आईचे नाव'][x]
##            =df['राष्ट्रीयत्व'][x]
##            =df['मातृभाषा'][x]
##            =df['धर्म'][x]
##            =df['जात'][x]
##            =df['पोटजात'][x]
##            =df['जन्मस्थळ'][x]
##            =df['तालुका'][x]
##            =df['जिल्हा'][x]
##            =df['राज्य'][x]
##            =df['देश'][x]
##            =df['इ.सनाप्रमाणे जन्मदिनांक'][x]
##            =df['जन्मदिनांक अक्षरी'][x]
##            =df['या पूर्वीची शाळा व इयत्ता '][x]
##            =df['या शाळेत प्रवेश घेतल्याचा दिनांक '][x]
##            =df['इयत्ता '][x]
##            =df['अभ्यासातली प्रगती'][x]
##            =df['वर्तणूक '][x]
##            =df['शाळा सोडल्याचा दिनांक'][x]
##            =df['कोणत्या इयत्तेत शिकत होता व केव्हापासून'][x]
##            =df['शाळा सोडण्याचे कारण '][x]
##            =df['शेरा'][x]

    gr=l[0]
    gr.text="\nअनु. क्र " + dat['अ. क्र. ']
    gr.add_run(str("जनरल रजि. क्र. " + dat['जनरल रजिस्टर क्र.']).rjust(126))  ##Gr. no. and count

    s_id=l[3]
    s_id.text = s_id.text +" "+ dat['स्टुडंट आय डी']  ##Student id

    uid=l[4]
    uid.text = uid.text + " " + dat['यु आय डी नं.'] ##uid no.

##    name=l[5]
##    name.text = name.text + " " + dat['विद्यार्थ्याचे संपूर्ण नाव'] ##name

    t[0].rows[0].cells[0].text=dat['नाव ']
    t[0].rows[0].cells[1].text=dat['वडिलांचे नाव ']
    t[0].rows[0].cells[2].text=dat['आडनाव']

    m_name=l[7]
    print(m_name.text+"abcd")
    m_name.text = "आईचे नाव :" + " " + dat['आईचे नाव'] ##mothers name

    nationality=l[8]
    nationality.text="राष्ट्रीयत्व : "+dat['राष्ट्रीयत्व']+"\t"
    nationality.add_run("मातृभाषा : " + dat['मातृभाषा']) ##nationality and mothertounge

    rel=l[9]
    rel.text="धर्म : "+dat['धर्म']+"\t"
    rel.add_run("जात : "+dat['जात']+"\t")
    rel.add_run("पोटजात : "+dat['पोटजात']+"\t") ##religion caste sub caste

    birthplace=l[10]
    birthplace.text="जन्मस्थळ (गांव/शहर) : "+dat['जन्मस्थळ']+"\t"
    birthplace.add_run("तालुका : "+dat['तालुका']+"\t")
    birthplace.add_run("जिल्हा : "+dat['जिल्हा']) ##birthplace village sub district district

    state=l[11]
    state.text="राज्य : "+dat['राज्य']+"\t"
    state.add_run("देश : "+dat['देश'])  ##state country

    bday=l[12]
    bday.text = bday.text + " " + dat['इ.सनाप्रमाणे जन्मदिनांक'] ##Birthdate

    bdayw=l[13]
    bdayw.text = bdayw.text + " " + dat['जन्मदिनांक अक्षरी'] ##Birthdate in words

    prev_sch=l[14]
    prev_sch.text = prev_sch.text + " " + dat['या पूर्वीची शाळा व इयत्ता '] ##Previous school and standard

    do_join=l[15]
    do_join.text="या शाळेत प्रवेश घेतल्याचा दिनांक : "+dat['या शाळेत प्रवेश घेतल्याचा दिनांक ']+"\t"
    do_join.add_run("इयत्ता : "+dat['इयत्ता '])  ##Date of join and standard

    prog=l[16]
    prog.text="अभ्यासातील प्रगती : "+dat['अभ्यासातली प्रगती']+"\t"
    prog.add_run("वर्तणूक : "+dat['वर्तणूक '])  ##Progress and remark

    do_leave=l[17]
    do_leave.text = do_leave.text + " " + dat['शाळा सोडल्याचा दिनांक'] ##date of leaving

    standard=l[18]
    standard.text = standard.text + " " + dat['कोणत्या इयत्तेत शिकत होता व केव्हापासून'] ##standard and since when

    reason=l[19]
    reason.text = reason.text + " " + dat['शाळा सोडण्याचे कारण '] ##reason for leaving

    remark=l[20]
    remark.text = remark.text + " " + dat['शेरा'] ##remark

    #print(l[21].text)

    ##for x in l:
    ##    print(x.text)

    ##
    ##sections = d.sections
    ##for section in sections:
    ##    section.top_margin = 10
    ##    section.bottom_margin = 10
    ##    section.left_margin = 10
    ##    section.right_margin = 10

    #d.sections[0].header.top_margin=10

    ##p=d.sections[0].header.paragraphs[0]
    ##p.text="""छत्रपती शिक्षण मंडळ, कल्याण
    ##माध्यमिक विद्यालय
    ##विवेकानंद संकुल, सानपाडा-नवी मुंबई, ठाणे फोन नं. ०२२-२७७५३७७७
    ##E-mail: mvvs1615038sanpada@rediffmail.com
    ##(जा.क्र.शिउसं/माध्य.-२/प्रमा/९८-९९/१०६५/५८)
    ##(शिक्षण उपसंचालक नशिक विभाग, नशिक यांचे कार्यालय दि. १५-३-९९)
    ##माध्यम : मराठी"""
    ##p.alignment=1
    ##p.line_spacing=1
    ##
    ##p1=d.add_paragraph()
    ##p1.add_run("\nअनु. क्र ")
    ##p1.add_run("जनरल रजि. क्र. ".rjust(126))
    ##p1.add_run("\n"+"{0:_^105}".format(""))
    ##p2=d.add_paragraph("यू-डायस क्र. 27211006404")
    ##p2.add_run("बोर्ड : मुंबई".center(74))
    ##p2.add_run("संलग्नता क्रमांक. 16.15.038")
    ##p3=d.add_paragraph()
    ##p3.alingnment=WD_ALIGN_PARAGRAPH.CENTER
    ##p3.bold=True
    ##p3.add_run("शाळा सोडल्याचे प्रमाण पत्र").bold=True
    ##p3.alignment=1
    ##p4=d.add_paragraph("स्टुडन्ट आय डी :")
    ##p5=d.add_paragraph("यू  आय डी नं :")
    ##p6=d.add_paragraph("विद्यार्थ्याचे संपूर्ण नाव :")
    ##p7=d.add_paragraph("आईचे नाव :")
    ##p8=d.add_paragraph("राष्ट्रीयत्व :"+"\t")
    ##p8.add_run("मातृभाषा :")
    ##p9=d.add_paragraph("धर्म :"+"\t")
    ##p9.add_run("जात :"+"\t")
    ##p9.add_run("पोटजात :"+"\t")

    ##p12=d.add_paragraph("इ. सनाप्रमाणे जन्मदिनांक :")
    ##p13=d.add_paragraph("जन्मदिनांक अक्षरी :")
    ##p14=d.add_paragraph("या पूर्वीची शाळा व इयत्ता :")
    ##p15=d.add_paragraph("या शाळेत प्रवेश घेतल्याचा दिनांक :"+"\t")
    ##p15.add_run("इयत्ता :")
    ##p17=d.add_paragraph("शाळा सोडल्याचा दिनांक :")
    ##p18=d.add_paragraph("कोणत्या इयत्तेत शिकत होता व केव्हापासून (अक्षरी व अंकी):")
    ##p19=d.add_paragraph("शाळा सोडण्याचे कारण :")
    ##p20=d.add_paragraph("शेरा :")
    ##p21=d.add_paragraph("दाखला देण्यात येतो की, वरील माहिती शाळेतील जनरल रजिस्टर नोंदी प्रमाणे आहे.")
    ###p22=d.add_paragraph("दिनांक :")
    ##
    ##pf=d.sections[0].footer.paragraphs[0]
    ##pf.text="""दिनांक :........."""

    filename='op.docx'
    d.save(filename)
    ##s="माध्यमिक विद्यालय"
    ##f=open(filename, "w",encoding='utf-8').write("{0:^175}".format("छत्रपती शिक्षण मंडळ, कल्याण")+
    ##                                             "\n"+
    ##                                             "{0:=^80}".format("माध्यमिक विद्यालय")+
    ##                                             "\n"+
    ##                                             str.center(s, 80,'.'))#\nविवेकानंद संकुल, सानपाडा-नवी मुंबई, ठाणे फोन नं. ०२२-२७७५३७७७\nE-mail: mvvs1615038sanpada@rediffmail.com\nमाध्यम : मराठी")))
    #f.write("helloworld\n")
    ##for x in range(0,101):
    ##    f.write("a")
    ##for x in f:
    ##    f1.write("{0:^40}".format(str(x)))
    #f1.close()


    win32api.ShellExecute (
      0,
      "print",
      filename,
      #
      # If this is None, the default printer will
      # be used anyway.
      #
     '/d:"%s"' % win32print.GetDefaultPrinter (),
      ".",
      0
    )

##import win32print
##import win32ui
##from PIL import Image, ImageWin
##
###
### Constants for GetDeviceCaps
###
###
### HORZRES / VERTRES = printable area
###
##HORZRES = 8
##VERTRES = 10
###
### LOGPIXELS = dots per inch
###
##LOGPIXELSX = 88
##LOGPIXELSY = 90
###
### PHYSICALWIDTH/HEIGHT = total area
###
##PHYSICALWIDTH = 110
##PHYSICALHEIGHT = 111
###
### PHYSICALOFFSETX/Y = left / top margin
###
##PHYSICALOFFSETX = 112
##PHYSICALOFFSETY = 113
##
##printer_name = win32print.GetDefaultPrinter ()
##file_name = "test.jpg"
##
###
### You can only write a Device-independent bitmap
###  directly to a Windows device context; therefore
###  we need (for ease) to use the Python Imaging
###  Library to manipulate the image.
###
### Create a device context from a named printer
###  and assess the printable size of the paper.
###
##hDC = win32ui.CreateDC ()
##hDC.CreatePrinterDC (printer_name)
##printable_area = hDC.GetDeviceCaps (HORZRES), hDC.GetDeviceCaps (VERTRES)
##printer_size = hDC.GetDeviceCaps (PHYSICALWIDTH), hDC.GetDeviceCaps (PHYSICALHEIGHT)
##printer_margins = hDC.GetDeviceCaps (PHYSICALOFFSETX), hDC.GetDeviceCaps (PHYSICALOFFSETY)
##
###
### Open the image, rotate it if it's wider than
###  it is high, and work out how much to multiply
###  each pixel by to get it as big as possible on
###  the page without distorting.
###
##bmp = Image.open (file_name)
##if bmp.size[0] > bmp.size[1]:
##  bmp = bmp.rotate (90)
##
##ratios = [1.0 * printable_area[0] / bmp.size[0], 1.0 * printable_area[1] / bmp.size[1]]
##scale = min (ratios)
##
###
### Start the print job, and draw the bitmap to
###  the printer device at the scaled size.
###
##hDC.StartDoc (file_name)
##hDC.StartPage ()
##
##dib = ImageWin.Dib (bmp)
##scaled_width, scaled_height = [int (scale * i) for i in bmp.size]
##x1 = int ((printer_size[0] - scaled_width) / 2)
##y1 = int ((printer_size[1] - scaled_height) / 2)
##x2 = x1 + scaled_width
##y2 = y1 + scaled_height
##dib.draw (hDC.GetHandleOutput (), (x1, y1, x2, y2))
##
##hDC.EndPage ()
##hDC.EndDoc ()
##hDC.DeleteDC ()
