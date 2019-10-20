import pandas
#from lc1 import give_print
from pandas import *
df=read_csv('dbaa.csv')
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from docx import *
import tempfile
import win32api
import win32print

def print_all():
    #merged_document = Document()
    filename = 'lc - Copy.docx'
    d1=docx.Document(filename)
    d1.add_page_break()

    for x in range(0,len(df['जनरल रजिस्टर क्र.'])):
        print(df['आडनाव'][x])
        #print(e.get())
        dat={}
        dat['अ. क्र. ']=str(int(df['अ. क्र. '][x]))
        dat['जनरल रजिस्टर क्र.']=str(df['जनरल रजिस्टर क्र.'][x])
        dat['स्टुडंट आय डी']=str(df['स्टुडंट आय डी'][x])
        dat['यु आय डी नं.']=str(int(df['यु आय डी नं.'][x]))
        print(str(int(df['यु आय डी नं.'][x])))
        dat['आडनाव']=str(df['आडनाव'][x])
        dat['नाव ']=str(df['नाव '][x])
        dat['वडिलांचे नाव ']=str(df['वडिलांचे नाव '][x])
        dat['आईचे नाव']=str(df['आईचे नाव'][x])
        dat['राष्ट्रीयत्व']=str(df['राष्ट्रीयत्व'][x])
        dat['मातृभाषा']=str(df['मातृभाषा'][x])
        dat['धर्म']=str(df['धर्म'][x])
        dat['जात']=str(df['जात'][x])
        dat['पोटजात']=str(df['पोटजात'][x])
        dat['जन्मस्थळ']=str(df['जन्मस्थळ'][x])
        dat['तालुका']=str(df['तालुका'][x])
        dat['जिल्हा']=str(df['जिल्हा'][x])
        dat['राज्य']=str(df['राज्य'][x])
        dat['देश']=str(df['देश'][x])
        dat['इ.सनाप्रमाणे जन्मदिनांक']=str(df['इ.सनाप्रमाणे जन्मदिनांक'][x])
        dat['जन्मदिनांक अक्षरी']=str(df['जन्मदिनांक अक्षरी'][x])
        dat['या पूर्वीची शाळा व इयत्ता ']=str(df['या पूर्वीची शाळा व इयत्ता '][x])
        dat['या शाळेत प्रवेश घेतल्याचा दिनांक ']=str(df['या शाळेत प्रवेश घेतल्याचा दिनांक '][x])
        dat['इयत्ता ']=str(df['इयत्ता '][x])
        dat['अभ्यासातली प्रगती']=str(df['अभ्यासातली प्रगती'][x])
        dat['वर्तणूक ']=str(df['वर्तणूक '][x])
        dat['शाळा सोडल्याचा दिनांक']=str(df['शाळा सोडल्याचा दिनांक'][x])
        dat['कोणत्या इयत्तेत शिकत होता व केव्हापासून']=str(df['कोणत्या इयत्तेत शिकत होता व केव्हापासून'][x])
        dat['शाळा सोडण्याचे कारण ']=str(df['शाळा सोडण्याचे कारण '][x])
        dat['शेरा']=str(df['शेरा'][x])
        #print(dat)

        
        #f=open(filename,'x')
        filename = 'lc - Copy.docx'
        d=docx.Document(filename)
        l=d.paragraphs
        t=d.tables


        gr=l[0]
        gr.text="\nअनु. क्र " + dat['अ. क्र. ']
        gr.add_run(str("जनरल रजि. क्र. " + dat['जनरल रजिस्टर क्र.']).rjust(126))  ##Gr. no. and count

        s_id=l[3]
        s_id.text = s_id.text +" "+ dat['स्टुडंट आय डी']  ##Student id

        uid=l[4]
        uid.text = uid.text + " " + dat['यु आय डी नं.'] ##uid no.

        t[0].rows[0].cells[0].text=dat['नाव ']
        t[0].rows[0].cells[1].text=dat['वडिलांचे नाव ']
        t[0].rows[0].cells[2].text=dat['आडनाव']

        m_name=l[7]
        print(m_name.text+"abcd")
        m_name.text = "आईचे नाव :" + " " + dat['आईचे नाव'] ##mothers name

        nationality=l[8]
        nationality.text="राष्ट्रीयत्व : "+dat['राष्ट्रीयत्व']+"\t"
        nationality.add_run("मातृभाषा : " + dat['मातृभाषा']) ##nationality and mothertounge

        rel=l[9]
        rel.text="धर्म : "+dat['धर्म']+"\t"
        rel.add_run("जात : "+dat['जात']+"\t")
        rel.add_run("पोटजात : "+dat['पोटजात']+"\t") ##religion caste sub caste

        birthplace=l[10]
        birthplace.text="जन्मस्थळ (गांव/शहर) : "+dat['जन्मस्थळ']+"\t"
        birthplace.add_run("तालुका : "+dat['तालुका']+"\t")
        birthplace.add_run("जिल्हा : "+dat['जिल्हा']) ##birthplace village sub district district

        state=l[11]
        state.text="राज्य : "+dat['राज्य']+"\t"
        state.add_run("देश : "+dat['देश'])  ##state country

        bday=l[12]
        bday.text = bday.text + " " + dat['इ.सनाप्रमाणे जन्मदिनांक'] ##Birthdate

        bdayw=l[13]
        bdayw.text = bdayw.text + " " + dat['जन्मदिनांक अक्षरी'] ##Birthdate in words

        prev_sch=l[14]
        prev_sch.text = prev_sch.text + " " + dat['या पूर्वीची शाळा व इयत्ता '] ##Previous school and standard

        do_join=l[15]
        do_join.text="या शाळेत प्रवेश घेतल्याचा दिनांक : "+dat['या शाळेत प्रवेश घेतल्याचा दिनांक ']+"\t"
        do_join.add_run("इयत्ता : "+dat['इयत्ता '])  ##Date of join and standard

        prog=l[16]
        prog.text="अभ्यासातील प्रगती : "+dat['अभ्यासातली प्रगती']+"\t"
        prog.add_run("वर्तणूक : "+dat['वर्तणूक '])  ##Progress and remark

        do_leave=l[17]
        do_leave.text = do_leave.text + " " + dat['शाळा सोडल्याचा दिनांक'] ##date of leaving

        standard=l[18]
        standard.text = standard.text + " " + dat['कोणत्या इयत्तेत शिकत होता व केव्हापासून'] ##standard and since when

        reason=l[19]
        reason.text = reason.text + " " + dat['शाळा सोडण्याचे कारण '] ##reason for leaving

        remark=l[20]
        remark.text = remark.text + " " + dat['शेरा'] ##remark

        for element in d.element.body:
                d1.element.body.append(element)

        
    d1.save("op.docx")

    win32api.ShellExecute (
          0,
          "print",
          "op.docx",
          #
          # If this is None, the default printer will
          # be used anyway.
          #
         '/d:"%s"' % win32print.GetDefaultPrinter (),
          ".",
          0
        )
